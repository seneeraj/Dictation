import streamlit as st
from pydub import AudioSegment
from vosk import Model, KaldiRecognizer
import wave
import json
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import tempfile

# Load Vosk Hindi model (must be downloaded and unzipped beforehand)
@st.cache_resource
def load_model():
    return Model("models/vosk-model-small-hi-0.22")

model = load_model()

def transcribe_audio(audio_path):
    wf = wave.open(audio_path, "rb")
    rec = KaldiRecognizer(model, wf.getframerate())
    rec.SetWords(True)
    results = []

    while True:
        data = wf.readframes(4000)
        if len(data) == 0:
            break
        if rec.AcceptWaveform(data):
            part_result = json.loads(rec.Result())
            results.append(part_result.get("text", ""))
    final_result = json.loads(rec.FinalResult())
    results.append(final_result.get("text", ""))
    return " ".join(results)

def save_to_docx(text, output_path):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Mangal'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Mangal')
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(14)
    run.font.name = 'Mangal'
    document.save(output_path)

# Streamlit UI
st.set_page_config(page_title="🗣️ Hindi Audio Transcriber", layout="centered")
st.title("🎙️ Audio to Hindi Text (.docx) – Vosk + Streamlit")

uploaded_file = st.file_uploader("Upload audio file", type=["mp3", "wav", "m4a", "flac"])

if uploaded_file:
    # Convert to WAV (Vosk needs WAV PCM)
    audio = AudioSegment.from_file(uploaded_file)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as f:
        audio.export(f.name, format="wav")
        wav_path = f.name

    st.info("🔍 Transcribing... Please wait.")
    try:
        text = transcribe_audio(wav_path)
        os.remove(wav_path)
    except Exception as e:
        st.error(f"❌ Transcription failed: {e}")
        st.stop()

    edited_text = st.text_area("📝 Edit the transcription:", value=text, height=300)
    save_path = st.text_input("📁 Save file as (without .docx):", value="hindi_transcription")

    if st.button("💾 Save as .docx"):
        try:
            full_path = f"{save_path}.docx"
            save_to_docx(edited_text, full_path)
            st.success(f"✅ Saved at: {full_path}")
        except Exception as e:
            st.error(f"❌ Failed to save: {e}")
